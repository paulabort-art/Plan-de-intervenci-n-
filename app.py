{\rtf1\ansi\ansicpg1252\cocoartf2639
\cocoatextscaling0\cocoaplatform0{\fonttbl\f0\fswiss\fcharset0 Helvetica;\f1\fnil\fcharset0 AppleColorEmoji;\f2\fnil\fcharset77 ZapfDingbatsITC;
}
{\colortbl;\red255\green255\blue255;}
{\*\expandedcolortbl;;}
\paperw11900\paperh16840\margl1440\margr1440\vieww11520\viewh8400\viewkind0
\pard\tx566\tx1133\tx1700\tx2267\tx2834\tx3401\tx3968\tx4535\tx5102\tx5669\tx6236\tx6803\pardirnatural\partightenfactor0

\f0\fs24 \cf0 import streamlit as st\
import pandas as pd\
from PIL import Image\
from io import BytesIO\
from docx import Document\
from docx.shared import Inches, Pt\
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT\
from docx.oxml import OxmlElement\
from docx.oxml.ns import qn\
import datetime\
\
st.set_page_config(page_title="Plan Intervenci\'f3n v2.1a", layout="wide")\
\
st.title("Plan Personalizado de Intervenci\'f3n \'97 v2.1a")\
\
# --- Sidebar para subir archivos ---\
st.sidebar.header("Archivos")\
uploaded_csv = st.sidebar.file_uploader("Sube un CSV (una o varias columnas)", type=["csv"])\
uploaded_logo = st.sidebar.file_uploader("Sube tu logo (opcional)", type=["png", "jpg", "jpeg"])\
\
# Logo por defecto\
if uploaded_logo:\
    logo = Image.open(uploaded_logo)\
else:\
    try:\
        logo = Image.open("logo.png")\
    except:\
        logo = None\
\
# --- Curso acad\'e9mico autom\'e1tico + editable ---\
now = datetime.datetime.now()\
curso_predeterminado = f"\{now.year\}-\{now.year+1\}"\
curso_acad = st.text_input("Curso acad\'e9mico", value=curso_predeterminado)\
\
# --- Encabezado: logo + curso ---\
col1, col2, col3 = st.columns([1,1,1])\
with col1:\
    if logo is not None:\
        st.image(logo, width=220)\
with col3:\
    st.markdown(f"**Curso acad\'e9mico:** \{curso_acad\}")\
\
st.markdown(\
    f"<div style='text-align:center; font-family:Calibri, sans-serif;'>"\
    f"<h2>PLAN PERSONALIZADO DE INTERVENCI\'d3N \'96 Aula de Audici\'f3n y Lenguaje (Versi\'f3n \{curso_acad\})</h2>"\
    f"</div>",\
    unsafe_allow_html=True\
)\
\
# --- Formulario de datos del alumno ---\
with st.form("datos_alumno"):\
    nombre = st.text_input("ALUMNO/A")\
    fecha_nac = st.date_input("Fecha de nacimiento", value=None)\
    curso = st.text_input("CURSO")\
    tutor = st.text_input("TUTOR/A")\
    maestra_al = st.text_input("MAESTRA AL", value="Paula Bort Museros")\
    diagnostico = st.text_area("DIAGN\'d3STICO", height=80)\
    fecha_inicio = st.date_input("FECHA DE INICIO", value=datetime.date.today())\
    numero_sesiones = st.text_input("N\'daMERO DE SESIONES")\
    situacion_actual = st.text_area("SITUACI\'d3N ACTUAL", height=100)\
    atencion = st.text_area("ATENCI\'d3N EDUCATIVA ESPEC\'cdFICA", height=100)\
    observaciones = st.text_area("OBSERVACIONES", height=100)\
    obs_trim1 = st.text_area("Observaciones 1\'ba Trimestre", height=60)\
    obs_trim2 = st.text_area("Observaciones 2\'ba Trimestre", height=60)\
    obs_trim3 = st.text_area("Observaciones 3\'ba Trimestre", height=60)\
    submitted = st.form_submit_button("Guardar datos")\
\
# --- Lectura del CSV ---\
if uploaded_csv is not None:\
    df_in = pd.read_csv(uploaded_csv)\
    if df_in.shape[1] == 1:\
        df_in.columns = ["Objetivo"]\
    else:\
        df_in = df_in.iloc[:, [0]].rename(columns=\{df_in.columns[0]: "Objetivo"\})\
    st.success("CSV cargado correctamente 
\f1 \uc0\u9989 
\f0 ")\
else:\
    try:\
        df_in = pd.read_csv("ejemplo.csv")\
        st.info("Usando el CSV de ejemplo (no se ha subido ninguno).")\
    except:\
        df_in = pd.DataFrame(\{"Objetivo": []\})\
        st.warning("No se encontr\'f3 CSV. Puedes subir uno o usar un ejemplo.")\
\
# --- Tabla editable de objetivos ---\
objetivos = df_in["Objetivo"].astype(str).tolist()\
bloques = ["SIN INICIAR", "NECESITA MEJORAR", "EST\'c1 PROGRESANDO", "CONSEGUIDO"]\
trimestres = ["1\'baT", "2\'baT", "3\'baT"]\
cols = ["Objetivo"] + [f"\{b\} \{t\}" for b in bloques for t in trimestres]\
data = [[o] + [""] * (len(cols) - 1) for o in objetivos]\
df_editable = pd.DataFrame(data, columns=cols)\
\
st.subheader("Tabla editable de objetivos")\
tabla = st.data_editor(df_editable, num_rows="dynamic", use_container_width=True)\
\
# --- Vista previa con colores desaturados ---\
if st.button("Vista previa de la tabla"):\
    colores = \{\
        "SIN INICIAR": "#EDEDED",\
        "NECESITA MEJORAR": "#FFE7C2",\
        "EST\'c1 PROGRESANDO": "#FFFCD2",\
        "CONSEGUIDO": "#DFF4D2"\
    \}\
    html = "<div style='font-family:Calibri;font-size:11pt;overflow:auto;'><table style='border-collapse:collapse;width:100%;'>"\
    html += "<tr><th rowspan='2' style='border:1px solid #000;padding:6px;width:40%;'>OBJETIVO</th>"\
    for b in bloques:\
        html += f"<th colspan='3' style='border:1px solid #000;padding:6px;background:\{colores[b]\};'>\{b\}</th>"\
    html += "</tr><tr>"\
    for _ in bloques:\
        for t in trimestres:\
            html += f"<th style='border:1px solid #000;padding:4px;width:6%;'>\{t\}</th>"\
    html += "</tr>"\
    for _, fila in tabla.iterrows():\
        html += "<tr>"\
        html += f"<td style='border:1px solid #000;padding:6px;'>\{fila['Objetivo']\}</td>"\
        for b in bloques:\
            for t in trimestres:\
                valor = fila[f"\{b\} \{t\}"]\
                texto = "X" if str(valor).strip().lower() in ["x","
\f1 \uc0\u10060 
\f0 ","
\f2 \uc0\u10006 
\f0 ","
\f2 \uc0\u10005 
\f0 "] else ("" if str(valor).strip()=="" else str(valor))\
                html += f"<td style='border:1px solid #000;text-align:center;background:\{colores[b]\};'>\{texto\}</td>"\
        html += "</tr>"\
    html += "</table></div>"\
    st.markdown(html, unsafe_allow_html=True)\
\
# --- Firma y generaci\'f3n del Word ---\
firma = st.text_input("Texto firma", value="Paula Bort Museros\\nMaestra de Audici\'f3n y Lenguaje")\
if st.button("Generar Word"):\
    doc = Document()\
    estilo = doc.styles["Normal"]\
    estilo.font.name = "Calibri"\
    estilo.font.size = Pt(11)\
\
    # Encabezado\
    tabla_encabezado = doc.add_table(rows=1, cols=2)\
    if logo:\
        bio = BytesIO()\
        logo.save(bio, format="PNG")\
        bio.seek(0)\
        tabla_encabezado.cell(0,0).paragraphs[0].add_run().add_picture(bio, width=Inches(1.9))\
    tabla_encabezado.cell(0,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT\
    tabla_encabezado.cell(0,1).paragraphs[0].add_run(f"Curso acad\'e9mico: \{curso_acad\}")\
\
    # T\'edtulo\
    p_titulo = doc.add_paragraph()\
    p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER\
    p_titulo.add_run(f"PLAN PERSONALIZADO DE INTERVENCI\'d3N \'96 Aula de Audici\'f3n y Lenguaje (Versi\'f3n \{curso_acad\})").bold = True\
\
    # Apartados principales\
    doc.add_paragraph("1. DATOS DEL ALUMNO").runs[0].bold = True\
    info = doc.add_table(rows=7, cols=2)\
    etiquetas = ["ALUMNO/A", "FECHA DE NACIMIENTO", "CURSO", "TUTOR/A", "MAESTRA AL", "DIAGN\'d3STICO", "FECHA DE INICIO"]\
    valores = [nombre, fecha_nac.strftime("%d/%m/%Y") if fecha_nac else "", curso, tutor, maestra_al, diagnostico, fecha_inicio.strftime("%d/%m/%Y")]\
    for i, et in enumerate(etiquetas):\
        info.cell(i,0).text = et + ":"\
        info.cell(i,1).text = valores[i]\
\
    # Generar el documento\
    nombre_seguro = nombre.replace(" ", "_") if nombre.strip() else "sin_nombre"\
    nombre_archivo = f"Plan_Intervenci\'f3n_\{nombre_seguro\}_\{curso_acad\}.docx"\
    doc.save(nombre_archivo)\
    with open(nombre_archivo, "rb") as f:\
        st.download_button("Descargar documento Word", data=f, file_name=nombre_archivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")\
    st.success(f"Documento generado correctamente 
\f1 \uc0\u9989 
\f0  \{nombre_archivo\}")\
}